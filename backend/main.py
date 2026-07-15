# パッケージ
from contextlib import asynccontextmanager

from fastapi import Depends, FastAPI, HTTPException, Response
from pydantic import BaseModel
import requests
from docx import Document
from openpyxl import load_workbook
from datetime import datetime, timezone
from io import BytesIO
import os
import re
import asyncio
import json
import traceback
from functools import lru_cache
from config import (
    GOOGLE_TRANSLATE_API_KEY, GOOGLE_MAPS_API_KEY, SLACK_WEBHOOK_URL, SLACK_BOT_TOKEN, SLACK_USER_ID,
    GOOGLE_DRIVE_CREDENTIALS_PATH, GOOGLE_DRIVE_FOLDER_ID, GOOGLE_SHEETS_SPREADSHEET_ID,
    TEMPLATE_DIR,
    GOOGLE_DRIVE_OAUTH_CLIENT_ID, GOOGLE_DRIVE_OAUTH_CLIENT_SECRET, GOOGLE_DRIVE_OAUTH_REFRESH_TOKEN,
    GOOGLE_DRIVE_OAUTH_TOKEN_URI, GOOGLE_DRIVE_OAUTH_SCOPES,
)
from fastapi.middleware.cors import CORSMiddleware
from sqlalchemy.orm import Session

from auth_routes import router as auth_router
from database import get_db, init_db
from deps import get_current_user
from models import Application, ApplicationStatus, User

import uvicorn
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload
from googleapiclient.errors import HttpError
import httplib2
from google_auth_httplib2 import AuthorizedHttp
from google.oauth2.credentials import Credentials as UserCredentials
from googletrans import Translator
import romkan2
import config

# TEMPLATE_DIR は config から取得（Render では backend のカレントディレクトリ）

def _google_api_traceback_enabled() -> bool:
    """GOOGLE_API_TRACEBACK=1 / true / yes / on のときだけ Google API 周りでスタックトレースを出す。"""
    v = os.environ.get("GOOGLE_API_TRACEBACK", "").strip().lower()
    return v in ("1", "true", "yes", "on")


def _log_google_exception(context: str, exc: BaseException) -> None:
    if not _google_api_traceback_enabled():
        return
    print(f"--- Google API traceback ({context}) ---")
    traceback.print_exception(type(exc), exc, exc.__traceback__)

# SSL検証回避
requests.packages.urllib3.disable_warnings()

@asynccontextmanager
async def lifespan(app: FastAPI):
    init_db()
    yield


app = FastAPI(lifespan=lifespan)
app.include_router(auth_router)

# CORS 設定を追加
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # 必要なら特定のオリジンに制限可能
    allow_credentials=True,
    allow_methods=["*"],  # すべてのHTTPメソッドを許可（GET, POST, OPTIONSなど）
    allow_headers=["*"],  # すべてのヘッダーを許可
)

class FormData(BaseModel):
    companyName: str
    presidentName: str
    presidentNameLocal: str = ''
    presidentAddress: str
    presidentAddressLocal: str = ''
    birthyear: int
    birthmonth: int
    birthday: int
    purpose1: str
    purpose2: str = ''
    purpose3: str = ''
    purpose4: str = ''
    purpose5: str = ''
    email: str

# Google Maps API を使用して住所を日本語に変換
def get_japanese_address(address: str) -> str:
    params = {
        "address": address,
        "key": GOOGLE_MAPS_API_KEY,
        "language": "ja"
    }
    response = requests.get("https://maps.googleapis.com/maps/api/geocode/json", params=params, verify=False)
    geocode_result = response.json()

    if geocode_result.get("status") == "OK":
        if "〒" in geocode_result["results"][0]["formatted_address"]:
            return geocode_result["results"][0]["formatted_address"].split("〒")[1][8:]
        return geocode_result["results"][0]["formatted_address"]
    else:
        raise HTTPException(status_code=500, detail="Address conversion failed")


def get_japanese_address_katakana(address: str) -> str:
    """住所を日本語に変換し、ひらがな部分をカタカナにした表記で返す（(C社員住所)用）"""
    ja_address = get_japanese_address(address)
    return hiragana_to_katakana(ja_address)


def _roman_to_katakana_for_address(address: str) -> str:
    """住所文字列のローマ字部分をカタカナに変換（スペース・カンマで区切って各トークンを変換）"""
    if not address or not address.strip():
        return ""
    tokens = re.split(r"[\s,]+", address.strip())
    result = []
    for part in tokens:
        if not part:
            continue
        if part.isalpha():
            try:
                result.append(romkan2.to_katakana(part.lower()))
            except Exception:
                result.append(part)
        else:
            result.append(part)
    return " ".join(result)


def get_japanese_address_katakana_safe(address: str) -> str:
    """(C社員住所)用。Maps API で日本語→カタカナ。失敗時は入力住所をそのまま返す（無理やりカタカナにはしない）"""
    if not address or not address.strip():
        return ""
    try:
        result = get_japanese_address_katakana(address)
        if result and not _is_mostly_roman(result):
            return result
    except Exception:
        pass
    return address


def _openai_should_try(text: str) -> bool:
    """OpenAI を叩く前の簡易フィルタ（コスト削減＆不必要な呼び出し回避）"""
    if not text or not text.strip():
        return False
    # 英字がほとんどないなら、既存ロジックに任せる
    return bool(re.search(r"[A-Za-z]", text))


def _is_valid_openai_katakana_only(output: str) -> bool:
    if not output or not output.strip():
        return False

    out = output.strip()
    out = re.sub(r"\s+", " ", out)

    # Latin文字が混ざっていたら不採用
    if re.search(r"[A-Za-z]", out):
        return False
    # ひらがなが混ざっていたら不採用（要件: カタカナのみ）
    if re.search(r"[\u3040-\u309F]", out):
        return False

    # カタカナ（＋長音記号）を最低1文字含む
    if not re.search(r"[\u30A0-\u30FF\u30FC]", out):
        return False

    # 許可する文字セット（社名/住所で許容したい記号を追加）
    # - カタカナ、長音、数字、空白、記号（&,-,.,,、等）を許可
    if not re.fullmatch(
        r"[\u30A0-\u30FF\u30FC0-9\s&＆\-\.,、，・/’'\"“”「」『』（）()!\?？:；。、]{1,300}",
        out,
    ):
        return False

    return True


async def katakana_via_openai_only_words(text: str, kind: str) -> str:
    """
    OpenAI で「単語だけ（カタカナのみ / 語間スペースあり）」を生成する。
    失敗時は空文字を返す（呼び出し側でフォールバックする）。
    """
    api_key = os.environ.get("OPENAI_API_KEY", "").strip()
    if not api_key:
        return ""

    if not _openai_should_try(text):
        return ""

    model = os.environ.get("OPENAI_KATAKANA_MODEL", "gpt-4o-mini").strip() or "gpt-4o-mini"

    kind = kind.strip().lower()
    if kind == "company":
        instruction = (
            "Convert the given company name into natural Japanese katakana transcription. "
            "Return ONLY the katakana transcription with spaces between words. "
            "No explanations, no extra text, no quotes. "
            "Keep numbers unchanged. Keep symbols like '&' and '-' as they are (surround with spaces only if needed). "
            "Treat possessive 's (e.g., yuto's) as yutos (remove the apostrophe)."
        )
    elif kind == "name":
        instruction = (
            "Convert the given personal name into natural Japanese katakana transcription. "
            "Return ONLY the katakana transcription with spaces between words. "
            "No explanations, no extra text, no quotes. "
            "Keep numbers unchanged. Keep common separators (spaces)."
        )
    elif kind == "address":
        instruction = (
            "Convert the given address into Japanese katakana transcription. "
            "Return ONLY the katakana transcription with spaces between words. "
            "No explanations, no extra text, no quotes. "
            "Keep numbers unchanged. Preserve separators such as commas and hyphens (convert to Japanese-safe punctuation if needed)."
        )
    else:
        return ""

    url = "https://api.openai.com/v1/chat/completions"
    headers = {"Authorization": f"Bearer {api_key}"}
    payload = {
        "model": model,
        "temperature": 0,
        "messages": [
            {"role": "system", "content": "You output only the requested text."},
            {"role": "user", "content": f"{instruction}\n\nInput:\n{text}"},
        ],
    }

    def _call_openai():
        return requests.post(url, headers=headers, json=payload, timeout=60)

    try:
        resp = await asyncio.to_thread(_call_openai)
        if resp.status_code != 200:
            return ""
        data = resp.json()
        out = (
            data.get("choices", [{}])[0]
            .get("message", {})
            .get("content", "")
            .strip()
        )
        out = re.sub(r"[\r\n]+", " ", out).strip()
        out = re.sub(r"\s+", " ", out).strip()
        if _is_valid_openai_katakana_only(out):
            return out
        return ""
    except Exception:
        # OpenAI 呼び出しエラー時は黙ってフォールバックする
        return ""


async def get_address_katakana_for_documents(address: str) -> str:
    """
    文書用の住所カタカナ表記を取得する。
    1. まず Maps API ベースの get_japanese_address_katakana_safe を試す
       - 日本語住所が取得できた場合はそのカタカナ表記を利用
    2. それでも日本語になっていない／失敗した場合は、英語住所に対して
       translate_to_katakana_proper_noun を適用して固有名詞寄りのカタカナにする
    3. 最後まで失敗した場合は元の住所文字列を返す
    """
    if not address or not address.strip():
        return ""

    # OpenAI（先に試して、成立したらそれを優先）
    try:
        openai_katakana = await katakana_via_openai_only_words(address, kind="address")
        if openai_katakana:
            return openai_katakana
    except Exception:
        pass

    # 1. 既存の安全な Maps ベース変換
    base = get_japanese_address_katakana_safe(address)
    if base and not _is_mostly_roman(base):
        return base

    # 2. 固有名詞寄りのカタカナ変換を試す（例: \"123 Main Street, Tokyo\" → \"123 メイン ストリート、トウキョウ\"）
    try:
        katakana = await translate_to_katakana_proper_noun(address)
        if katakana:
            return katakana
    except Exception:
        pass

    # 3. すべて失敗した場合は元の住所を返す
    return address

# 翻訳関数
async def translate_text(text: str, target_lang: str = "ja") -> str:
    # 空文字列の場合は空文字列を返す（エラーを投げない）
    if not text or not text.strip():
        return ""
    
    max_retries = 3
    delay = 1
    
    for attempt in range(max_retries):
        try:
            # タイムアウト設定付きでTranslatorを作成（120秒）
            async with Translator(timeout=120.0) as translator:
                result = await translator.translate(text, dest=target_lang)
                return result.text
        except Exception as e:
            # 接続タイムアウトエラーの場合のみリトライ
            is_timeout = "ConnectTimeout" in str(type(e).__name__) or "timeout" in str(e).lower() or "10060" in str(e)
            if is_timeout and attempt < max_retries - 1:
                print(f"⚠️ 翻訳接続エラー（試行 {attempt + 1}/{max_retries}）、{delay}秒後にリトライ...")
                await asyncio.sleep(delay)
                delay *= 2  # 指数バックオフ
            else:
                # 最終的に失敗した場合、元のテキストを返す（フォールバック）
                if attempt == max_retries - 1:
                    print(f"⚠️ 翻訳失敗（{max_retries}回試行後）、元のテキストを使用: {text}")
                    return text
                else:
                    import traceback
                    error_detail = f"Translation failed: {str(e)}"
                    error_traceback = traceback.format_exc()
                    print(f"❌ {error_detail}")
                    print(f"❌ Traceback: {error_traceback}")
                    raise HTTPException(status_code=500, detail=error_detail)


async def translate_to_katakana_proper_noun(text: str) -> str:
    """
    固有名詞寄りのカタカナ表記を生成する。
    - 既に日本語を含む場合: ひらがなのみカタカナにして返す
    - 日本語を含まない場合: Google翻訳で日本語化 → ひらがなをカタカナに変換
      （結果がほぼローマ字の場合は不自然とみなし、フォールバックへ）
    - 翻訳に失敗／不自然な場合: romkan2 によるローマ字→カタカナ変換にフォールバック
    """
    if not text or not text.strip():
        return ""

    raw = text.strip()

    # 既に日本語が含まれている場合は、ひらがなをカタカナに揃えて返す
    if _contains_japanese(raw):
        return hiragana_to_katakana(raw)

    # 日本語を含まない場合は、まず翻訳 API を試す
    try:
        translated_ja = await translate_text(raw, target_lang="ja")
        if translated_ja:
            katakana = hiragana_to_katakana(translated_ja.replace(" ", ""))
            # 結果がほぼローマ字なら不自然なのでフォールバックに回す
            if not _is_mostly_roman(katakana):
                return katakana
    except Exception:
        # 翻訳 API が失敗した場合はフォールバックに移行
        pass

    # フォールバック: romkan2 によるローマ字→カタカナ変換
    return _roman_to_katakana(raw)


# カタカナ変換用関数（ひらがなに変換してからカタカナに）
async def to_katakana(text: str) -> str:
    # 空文字列の場合は空文字列を返す（エラーを投げない）
    if not text or not text.strip():
        return ""
    
    max_retries = 3
    delay = 1
    
    for attempt in range(max_retries):
        try:
            # タイムアウト設定付きでTranslatorを作成（120秒）
            async with Translator(timeout=120.0) as translator:
                # 日本語に翻訳（ひらがなになることが多い）
                result = await translator.translate(text, dest="ja")
                hiragana_text = result.text.replace(" ", "")
                # ひらがなをカタカナに変換
                return hiragana_to_katakana(hiragana_text)
        except Exception as e:
            # 接続タイムアウトエラーの場合のみリトライ
            is_timeout = "ConnectTimeout" in str(type(e).__name__) or "timeout" in str(e).lower() or "10060" in str(e)
            if is_timeout and attempt < max_retries - 1:
                print(f"⚠️ カタカナ変換接続エラー（試行 {attempt + 1}/{max_retries}）、{delay}秒後にリトライ...")
                await asyncio.sleep(delay)
                delay *= 2  # 指数バックオフ
            else:
                # 最終的に失敗した場合、元のテキストをカタカナ風に変換（フォールバック）
                if attempt == max_retries - 1:
                    print(f"⚠️ カタカナ変換失敗（{max_retries}回試行後）、元のテキストを使用: {text}")
                    # 元のテキストをそのまま返す（カタカナ変換なし）
                    return text
                else:
                    import traceback
                    error_detail = f"Katakana conversion failed: {str(e)}"
                    error_traceback = traceback.format_exc()
                    print(f"❌ {error_detail}")
                    print(f"❌ Traceback: {error_traceback}")
                    raise HTTPException(status_code=500, detail=error_detail)

# ひらがな→カタカナ変換関数
def hiragana_to_katakana(text: str) -> str:
    """ひらがなをカタカナに変換"""
    result = []
    for char in text:
        # ひらがなの範囲: U+3041-U+3096
        # カタカナの範囲: U+30A1-U+30F6
        # 差は0x60（96）
        if '\u3041' <= char <= '\u3096':
            # ひらがなをカタカナに変換
            katakana_char = chr(ord(char) + 0x60)
            result.append(katakana_char)
        else:
            # ひらがな以外はそのまま
            result.append(char)
    return ''.join(result)


def _is_mostly_roman(text: str) -> bool:
    """文字列の大半がローマ字（ASCII英字）かどうか。翻訳失敗で元のアルファベットが返った場合に True"""
    if not text or not text.strip():
        return True
    letters = [c for c in text if c.isalpha()]
    if not letters:
        return False
    roman = sum(1 for c in letters if ord(c) < 0x100 and c.isascii())
    return roman >= len(letters) * 0.5


def _contains_japanese(text: str) -> bool:
    """文字列にひらがな・カタカナ・漢字が含まれているか判定"""
    if not text:
        return False
    for ch in text:
        code = ord(ch)
        # ひらがな
        if 0x3040 <= code <= 0x309F:
            return True
        # カタカナ
        if 0x30A0 <= code <= 0x30FF:
            return True
        # CJK統合漢字
        if 0x4E00 <= code <= 0x9FFF:
            return True
    return False


def _is_japanese_token(token: str) -> bool:
    """トークン単位で日本語（ひらがな・カタカナ・漢字）を含むかどうか"""
    return _contains_japanese(token)


def _is_latin_token(token: str) -> bool:
    """トークンがローマ字のみ（A-Z, a-z と一部記号）かどうか"""
    if not token:
        return False
    t = token.strip()
    if not t:
        return False
    # アルファベットと一部の記号（' -）のみ許可
    return re.fullmatch(r"[A-Za-z][A-Za-z'’-]*", t) is not None


def _roman_to_katakana(roman_text: str) -> str:
    """ローマ字文字列をカタカナに変換（romkan2 使用）。スペースは区切りとして残す"""
    parts = roman_text.split()
    result = []
    for part in parts:
        part = part.strip()
        if not part:
            continue
        try:
            result.append(romkan2.to_katakana(part.lower()))
        except Exception:
            result.append(part)
    return " ".join(result)


def to_katakana_name(name: str) -> str:
    """氏名・商号などの固有名詞を翻訳せずにカタカナ表記にする"""
    if not name or not name.strip():
        return ""
    text = name.strip()
    # 既に日本語が含まれている場合は、ひらがなのみカタカナにして返す
    if _contains_japanese(text):
        return hiragana_to_katakana(text)
    # ローマ字などのラテン文字のみの場合は romkan2 でカタカナに変換
    return _roman_to_katakana(text)


def _has_ascii_letters(text: str) -> bool:
    """文字列中に ASCII 英字が含まれるかどうか"""
    if not text:
        return False
    return any(("A" <= c <= "Z") or ("a" <= c <= "z") for c in text)


async def roman_token_to_katakana_with_fallback(token: str) -> str:
    """
    ローマ字トークンをカタカナに変換する。
    1. まず romkan2 で変換を試みる
       - ASCII 英字が残らず、かつ元のトークンと明らかに異なる場合はその結果を採用
    2. romkan2 でうまく変換できなかった場合のみ、
       translate_to_katakana_proper_noun を使って固有名詞寄りのカタカナに変換する
       - 結果に ASCII 英字が含まれる場合は不自然とみなし、採用しない
    3. 最後までうまく変換できなければ元のトークンをそのまま返す
    """
    if not token:
        return ""

    t = token.strip()
    if not t:
        return ""

    # 1. romkan2 による機械的なローマ字→カタカナ変換
    try:
        romkan_result = romkan2.to_katakana(t.lower())
    except Exception:
        romkan_result = t

    if romkan_result and not _has_ascii_letters(romkan_result) and romkan_result.lower() != t.lower():
        return romkan_result

    # 2. romkan2 がうまく変換できなかった場合のみ、翻訳ベースのフォールバックを試す
    try:
        katakana = await translate_to_katakana_proper_noun(t)
        if katakana and not _has_ascii_letters(katakana):
            return katakana
    except Exception:
        # 翻訳側での失敗はログのみに留め、最終的には元のトークンにフォールバックする
        pass

    # 3. どうしても適切なカタカナが得られない場合は、元のトークンを返す
    return t


async def name_to_katakana_roman_only(name: str) -> str:
    """
    氏名を「意味翻訳せずに」ローマ字ベースでカタカナ表記にする。
    - 日本語（漢字・ひらがな・カタカナ）を含むトークンは、そのまま（ひらがなのみカタカナ化）
    - ローマ字のみのトークンは romkan2 を優先し、必要に応じて Google 翻訳にフォールバック
    - それ以外のトークンはそのまま
    例:
      \"Smirnov wachi\" -> 「スミルノフ ワチ」
      \"山田 wachi\" -> 「山田 ワチ」
    """
    if not name or not name.strip():
        return ""

    # OpenAI（先に試して、成立したらそれを優先）
    try:
        openai_katakana = await katakana_via_openai_only_words(name, kind="name")
        if openai_katakana:
            return openai_katakana
    except Exception:
        pass

    text = name.strip()
    tokens = re.split(r"\s+", text)
    converted: list[str] = []

    for token in tokens:
        if not token:
            continue
        if _is_japanese_token(token):
            # 日本語を含むトークンは、ひらがなのみカタカナに揃えて保持
            converted.append(hiragana_to_katakana(token))
        elif _is_latin_token(token):
            # ローマ字のみのトークンは romkan2 → 必要に応じて翻訳フォールバック
            converted_token = await roman_token_to_katakana_with_fallback(token)
            converted.append(converted_token)
        else:
            # 記号やその他の文字が混ざる場合は安全のためそのまま
            converted.append(token)

    return " ".join(converted)


async def name_to_katakana(name: str) -> str:
    """氏名を固有名詞として自然なカタカナに変換するヘルパー（translate_to_katakana_proper_noun の薄いラッパー）"""
    return await translate_to_katakana_proper_noun(name)


async def company_name_to_katakana(name: str) -> str:
    """
    会社名を固有名詞として自然なカタカナに変換する。
    例: 「yuto's food service&Co」→「ユトス フード サービス & コー」
    - 「's」は所有格として扱い、「yuto's」→「yutos」として変換
    - 「&」はそのまま残す
    - 「-」などのハイフンは区切り記号として扱い、出力からは基本的に除外
    - ローマ字トークンは romkan2 をベースにしつつ、必要に応じて翻訳フォールバックを利用
    - よくある英単語（food, service, Coなど）はカスタム辞書で自然なカタカナに
    """
    if not name or not name.strip():
        return ""

    # OpenAI（先に試して、成立したらそれを優先）
    try:
        openai_katakana = await katakana_via_openai_only_words(name, kind="company")
        if openai_katakana:
            return openai_katakana
    except Exception:
        pass

    # 記号の正規化
    text = name.strip().replace("’", "'")
    # & や - の前後に空白を入れてトークンとして分離
    text = re.sub(r"([&\-–—])", r" \1 ", text)
    # 空白でトークン分割（先に記号の前後へ空白を入れているため、これで安定する）
    raw_tokens = re.split(r"\s+", text.strip())

    # よく使う単語のカスタムマッピング（すべてカタカナ）
    custom_dict = {
        "food": "フード",
        "foods": "フーズ",
        "service": "サービス",
        "services": "サービス",
        "co": "コー",
        "co.": "コー",
        "company": "カンパニー",
        "inc": "インク",
        "inc.": "インク",
        "ltd": "リミテッド",
        "ltd.": "リミテッド",
        "llc": "エルエルシー",
        "corp": "コープ",
        "corp.": "コープ",
        "corporation": "コーポレーション",
        "holdings": "ホールディングス",
        "group": "グループ",
        "japan": "ジャパン",
        "tokyo": "トウキョウ",
    }

    converted: list[str] = []

    for token in raw_tokens:
        if not token:
            continue

        t = token.strip()
        if not t:
            continue

        # & はそのまま残す
        if t in {"&", "＆"}:
            converted.append("&")
            continue

        # ハイフン系は区切り記号として無視（出力には入れない）
        if t in {"-", "–", "—"}:
            continue

        # 日本語を含むトークンは、ひらがなのみカタカナに揃えて保持
        if _is_japanese_token(t):
            converted.append(hiragana_to_katakana(t))
            continue

        # 所有格 's をまとめて処理（yuto's → yutos）
        m = re.fullmatch(r"(.+)'s", t, flags=re.IGNORECASE)
        if m:
            base = m.group(1)
            t_norm = base + "s"
        else:
            t_norm = t

        # ピリオド付きの略称（Co. など）は末尾のピリオドを辞書検索のために一旦除去
        t_lower = t_norm.lower()
        if t_lower.endswith("."):
            t_lower = t_lower[:-1]

        # カスタム辞書にある単語は優先してそのカタカナを使う
        if t_lower in custom_dict:
            converted.append(custom_dict[t_lower])
            continue

        # ローマ字のみのトークンは、氏名と同様のロジックで変換
        if _is_latin_token(t_norm):
            converted_token = await roman_token_to_katakana_with_fallback(t_norm)
            converted.append(converted_token)
            continue

        # 上記いずれにも当てはまらない場合は、そのまま出力
        converted.append(t)

    # 空でないトークンをスペース区切りで結合
    return " ".join(token for token in converted if token)

def replace_in_docx_keeping_style(doc, replacements: dict):
    """docx内の全段落・表セルのテキストを置換する。run単位で置換し、残りは段落単位で置換（複数runに分かれているプレースホルダも確実に置換）"""
    for placeholder, value in replacements.items():
        if value is None:
            value = ""
        value = str(value)
        # 1) run単位で置換（書式保持）
        for p in doc.paragraphs:
            for run in p.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, value)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            if placeholder in run.text:
                                run.text = run.text.replace(placeholder, value)
        # 2) 段落全体でまだ残っていれば置換（複数runに分かれている場合のフォールバック）
        for p in doc.paragraphs:
            if placeholder in p.text:
                p.text = p.text.replace(placeholder, value)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if placeholder in p.text:
                            p.text = p.text.replace(placeholder, value)


# ファイル名を安全な形式に変換
def sanitize_filename(filename: str, max_length: int = 100) -> str:
    """ファイル名から特殊文字を削除し、安全な形式に変換"""
    # 特殊文字を削除または置換
    invalid_chars = ['/', '\\', ':', '*', '?', '"', '<', '>', '|']
    for char in invalid_chars:
        filename = filename.replace(char, '')
    # 空白をハイフンに置換
    filename = filename.replace(' ', '-')
    # 連続するハイフンを1つに
    while '--' in filename:
        filename = filename.replace('--', '-')
    # 先頭・末尾のハイフンを削除
    filename = filename.strip('-')
    # 最大長を制限
    if len(filename) > max_length:
        filename = filename[:max_length]
    return filename


def escape_drive_query_name(name: str) -> str:
    """Drive検索クエリ用にフォルダ名の特殊文字をエスケープ"""
    if not name:
        return ""
    # Drive q では、文字列リテラル内の `'` と `\` が特殊扱いになる
    # まずバックスラッシュをエスケープしてから、シングルクォートを `\'` としてエスケープする
    escaped = name.replace("\\", "\\\\")
    escaped = escaped.replace("'", "\\'")
    return escaped


# 出力ファイル名を生成
def generate_output_filename(company_name: str, file_type: str) -> str:
    """会社名と日付を含む出力ファイル名を生成"""
    date_str = datetime.now().strftime("%Y%m%d")
    safe_company_name = sanitize_filename(company_name)
    return f"{safe_company_name}-{date_str}-{file_type}"

def send_slack_notification(message: str):
    payload = {"text": message}
    try:
        response = requests.post(SLACK_WEBHOOK_URL, json=payload)
        response.raise_for_status()
    except Exception as e:
        print(f"Slack通知エラー: {e}")

# リンク送信
# def upload_file_to_slack(endpoint: str, title: str):
#     slack_api_url = "https://slack.com/api/chat.postMessage"
#     download_url = f"https://onestopjpn.onrender.com/{endpoint}"  # ←本番URLに変更してください

#     headers = {
#         "Authorization": f"Bearer {SLACK_BOT_TOKEN}",
#         "Content-Type": "application/json"
#     }

#     message = {
#         "channel": SLACK_USER_ID,
#         "text": f":white_check_mark: {title} を生成しました。\n📎 ダウンロード: <{download_url}>"
#     }

#     try:
#         response = requests.post(slack_api_url, headers=headers, json=message)
#         result = response.json()
#         print("Slack chat.postMessage response:", result)
#         if not result.get("ok"):
#             print(f"Slackメッセージ送信失敗: {result.get('error')}")
#     except Exception as e:
#         print("Slackメッセージ送信エラー:", e)

def upload_file_to_slack(filepath: str, title: str):
    print(f"📎 ファイルアップロード処理を開始：{filepath} → {SLACK_USER_ID}")
    with open(filepath, "rb") as file_content:
        response = requests.post(
            "https://slack.com/api/files.upload",
            headers={"Authorization": f"Bearer {SLACK_BOT_TOKEN}"},
            files={"file": (filepath, file_content)},
            data={
                "channels": SLACK_USER_ID,
                "title": title,
                "filename": filepath,
            }
        )
    res_json = response.json()
    if res_json.get("ok"):
        print(f"✅ ファイルアップロード成功（{title}）")
    else:
        print(f"❌ Slackファイルアップロード失敗: {res_json}")

@lru_cache(maxsize=1)
def get_google_drive_service():
    """Google Drive API を OAuth ユーザー（refresh_token）で取得。ファイルはユーザー所有でアップロードされる。"""
    try:
        if not (GOOGLE_DRIVE_OAUTH_CLIENT_ID and GOOGLE_DRIVE_OAUTH_CLIENT_SECRET and GOOGLE_DRIVE_OAUTH_REFRESH_TOKEN):
            raise RuntimeError(
                "Google Drive OAuth の環境変数が設定されていません。"
                "GOOGLE_DRIVE_OAUTH_CLIENT_ID / GOOGLE_DRIVE_OAUTH_CLIENT_SECRET / "
                "GOOGLE_DRIVE_OAUTH_REFRESH_TOKEN を設定してください。"
            )

        creds = UserCredentials(
            token=None,
            refresh_token=GOOGLE_DRIVE_OAUTH_REFRESH_TOKEN,
            token_uri=GOOGLE_DRIVE_OAUTH_TOKEN_URI,
            client_id=GOOGLE_DRIVE_OAUTH_CLIENT_ID,
            client_secret=GOOGLE_DRIVE_OAUTH_CLIENT_SECRET,
            scopes=GOOGLE_DRIVE_OAUTH_SCOPES,
        )

        http = httplib2.Http(timeout=300)
        authorized_http = AuthorizedHttp(creds, http=http)
        # num_retries を明示して一時的な通信失敗を吸収しやすくする
        return build(
            "drive",
            "v3",
            http=authorized_http,
            cache_discovery=False,
            num_retries=5,
        )
    except Exception as e:
        print(f"❌ Google Drive認証エラー: {e}")
        import traceback
        print(f"❌ Traceback: {traceback.format_exc()}")
        raise

def get_google_sheets_service():
    """Google Sheets APIサービスを取得"""
    try:
        # 認証情報ファイルの存在確認
        credentials_path = os.path.abspath(GOOGLE_DRIVE_CREDENTIALS_PATH)
        if not os.path.exists(credentials_path):
            raise FileNotFoundError(
                f"認証情報ファイルが見つかりません: {credentials_path}\n"
                f"現在の作業ディレクトリ: {os.getcwd()}\n"
                f"設定されたパス: {GOOGLE_DRIVE_CREDENTIALS_PATH}"
            )
        
        creds = service_account.Credentials.from_service_account_file(
            credentials_path,
            scopes=['https://www.googleapis.com/auth/spreadsheets']
        )
        
        # HTTPクライアントにタイムアウト設定を追加
        http = httplib2.Http(timeout=300)  # 5分
        authorized_http = AuthorizedHttp(creds, http=http)
        
        return build('sheets', 'v4', http=authorized_http)
    except Exception as e:
        print(f"❌ Google Sheets認証エラー: {e}")
        import traceback
        print(f"❌ Traceback: {traceback.format_exc()}")
        raise

async def create_company_folder(company_name: str) -> str:
    """会社名のフォルダをGoogleドライブに作成（重複チェック付き、リトライ付き）"""
    max_retries = 3
    delay = 2

    def _is_timeout_error(err: Exception) -> bool:
        # WinError 10060 / 通信タイムアウト系の判定
        s = str(err).lower()
        return (
            "10060" in s
            or "connecttimeout" in s
            or "timeout" in s
        )
    
    for attempt in range(max_retries):
        try:
            drive_service = get_google_drive_service()
            
            # フォルダ名を安全な形式に変換
            safe_folder_name = sanitize_filename(company_name)
            # Drive クエリ用にシングルクォートをエスケープした名前を使用
            escaped_name = escape_drive_query_name(safe_folder_name)
            
            # 既存のフォルダを検索（重複チェック）
            query = f"name='{escaped_name}' and mimeType='application/vnd.google-apps.folder' and trashed=false"
            if GOOGLE_DRIVE_FOLDER_ID:
                query += f" and '{GOOGLE_DRIVE_FOLDER_ID}' in parents"
            
            try:
                results = drive_service.files().list(
                    q=query,
                    fields='files(id, name)',
                    pageSize=1
                ).execute()
                
                items = results.get('files', [])
                if items:
                    # 既存のフォルダが見つかった場合はそのIDを返す
                    folder_id = items[0]['id']
                    print(f"✅ 既存のフォルダを使用: {safe_folder_name} (ID: {folder_id})")
                    return folder_id
            except Exception as search_error:
                # 検索エラーが発生した場合は、新規作成を試みる
                print(f"⚠️ フォルダ検索エラー（新規作成を試みます）: {search_error}")
                _log_google_exception("create_company_folder: files().list", search_error)
                # 検索段階でタイムアウトするなら、このリクエスト内ではフォルダ作成を諦める
                if _is_timeout_error(search_error):
                    print("⚠️ フォルダ検索がタイムアウトしたため、会社フォルダ作成をスキップします。親フォルダへ直接アップロードします。")
                    return ""
            
            # フォルダが存在しない場合は新規作成
            folder_metadata = {
                'name': safe_folder_name,
                'mimeType': 'application/vnd.google-apps.folder'
            }
            
            # 親フォルダIDが設定されている場合は指定
            if GOOGLE_DRIVE_FOLDER_ID:
                folder_metadata['parents'] = [GOOGLE_DRIVE_FOLDER_ID]
            
            # フォルダを作成
            folder = drive_service.files().create(
                body=folder_metadata,
                fields='id'
            ).execute()
            
            folder_id = folder.get('id')
            print(f"✅ Googleドライブフォルダ作成成功: {safe_folder_name} (ID: {folder_id})")
            return folder_id
        except HttpError as e:
            # Google Drive APIが有効化されていない場合のエラー
            if e.resp.status == 403 and 'accessNotConfigured' in str(e):
                error_msg = (
                    f"❌ Google Drive APIが有効化されていません。\n"
                    f"以下のURLからGoogle Drive APIを有効化してください：\n"
                    f"https://console.developers.google.com/apis/api/drive.googleapis.com/overview\n"
                    f"エラー詳細: {e}"
                )
                print(error_msg)
                raise Exception(error_msg) from e
            # その他のHttpError（権限エラーなど）
            if attempt < max_retries - 1:
                print(f"⚠️ Googleドライブ接続エラー（試行 {attempt + 1}/{max_retries}）、{delay}秒後にリトライ...")
                print(f"エラー詳細: {e}")
                await asyncio.sleep(delay)
                delay *= 2  # 指数バックオフ
            else:
                print(f"❌ Googleドライブ接続失敗（{max_retries}回試行後）: {e}")
                raise
        except (OSError, Exception) as e:
            # 接続タイムアウトなどのネットワークエラー
            is_timeout = "10060" in str(e) or "timeout" in str(e).lower() or "ConnectTimeout" in str(type(e).__name__)
            if is_timeout:
                # 会社フォルダ作成だけがコケても、ファイル自体は親フォルダへ直接アップロードできる
                print(f"⚠️ Google Drive がタイムアウトしたため、会社フォルダ作成をスキップします（親フォルダへ直接アップロード）: {e}")
                return ""

            if attempt < max_retries - 1:
                print(f"⚠️ Googleドライブ接続エラー（試行 {attempt + 1}/{max_retries}）、{delay}秒後にリトライ...")
                print(f"エラー詳細: {e}")
                await asyncio.sleep(delay)
                delay *= 2  # 指数バックオフ

            print(f"❌ Googleドライブ接続失敗（{max_retries}回試行後）: {e}")
            raise

async def upload_file_to_drive(file_path: str, folder_id: str, file_name: str = None) -> str:
    """ファイルをGoogleドライブにアップロード（リトライ付き）"""
    if not folder_id:
        raise ValueError("フォルダIDが指定されていません。GOOGLE_DRIVE_FOLDER_IDを確認してください。")

    max_retries = 3
    delay = 2
    
    for attempt in range(max_retries):
        try:
            drive_service = get_google_drive_service()
            if file_name is None:
                file_name = os.path.basename(file_path)
            
            # ファイルメタデータ
            file_metadata = {
                'name': file_name,
                'parents': [folder_id]
            }
            
            # ファイルをアップロード（共有ドライブ対応）
            media = MediaFileUpload(file_path, resumable=True)
            file = drive_service.files().create(
                body=file_metadata,
                media_body=media,
                fields='id',
                supportsAllDrives=True
            ).execute()
            
            file_id = file.get('id')
            print(f"✅ Googleドライブファイルアップロード成功: {file_name} (ID: {file_id})")
            return file_id
        except HttpError as e:
            # Google Drive APIが有効化されていない場合のエラー
            if e.resp.status == 403 and 'accessNotConfigured' in str(e):
                error_msg = (
                    f"❌ Google Drive APIが有効化されていません。\n"
                    f"以下のURLからGoogle Drive APIを有効化してください：\n"
                    f"https://console.developers.google.com/apis/api/drive.googleapis.com/overview\n"
                    f"エラー詳細: {e}"
                )
                print(error_msg)
                raise Exception(error_msg) from e
            # ストレージクォータ超過（OAuth 利用時は通常発生しない）
            if e.resp.status == 403 and 'storageQuotaExceeded' in str(e):
                error_msg = (
                    f"❌ Google ドライブのストレージ容量が不足しています。\n"
                    f"ドライブの空き容量を確認してください。\n"
                    f"エラー詳細: {e}"
                )
                print(error_msg)
                raise Exception(error_msg) from e
            # その他の 403（権限不足・アクセス拒否など）
            if e.resp.status == 403:
                print("❌ 403 権限エラー: フォルダIDや共有設定を確認してください。")
            # その他のHttpError（権限エラーなど）
            if attempt < max_retries - 1:
                print(f"⚠️ Googleドライブアップロード接続エラー（試行 {attempt + 1}/{max_retries}）、{delay}秒後にリトライ...")
                print(f"エラー詳細: {e}")
                await asyncio.sleep(delay)
                delay *= 2  # 指数バックオフ
            else:
                print(f"❌ Googleドライブアップロード接続失敗（{max_retries}回試行後）: {e}")
                _log_google_exception("upload_file_to_drive: HttpError", e)
                raise
        except (OSError, Exception) as e:
            # 接続タイムアウトなどのネットワークエラー
            is_timeout = "10060" in str(e) or "timeout" in str(e).lower() or "ConnectTimeout" in str(type(e).__name__)
            if is_timeout and attempt < max_retries - 1:
                print(f"⚠️ Googleドライブアップロード接続エラー（試行 {attempt + 1}/{max_retries}）、{delay}秒後にリトライ...")
                await asyncio.sleep(delay)
                delay *= 2  # 指数バックオフ
            else:
                print(f"❌ Googleドライブアップロード接続失敗（{max_retries}回試行後）: {e}")
                _log_google_exception("upload_file_to_drive: network", e)
                raise

def append_to_spreadsheet(data: FormData, file_paths: dict = None, folder_url: str = ""):
    """Spreadsheetに回答情報を追加（FolderURL 列がある場合はフォルダURLも記録）"""
    try:
        if not GOOGLE_SHEETS_SPREADSHEET_ID:
            print("⚠️ Spreadsheet IDが設定されていません。スキップします。")
            return
        
        sheets_service = get_google_sheets_service()
        
        # 現在の日時
        current_datetime = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        # 行データを準備（スプレッドシートの列順に合わせる）
        # CreatedDate, CompanyName,
        # RepresentativeName, RepresentativeNameLocal, RepresentativeBirthDay,
        # RepresentativeAddress, RepresentativeAddressLocal,
        # BusinessPurpose1-5, Email Address, FolderURL
        row_data = [
            current_datetime,  # CreatedDate
            data.companyName,  # CompanyName
            data.presidentName,  # RepresentativeName
            data.presidentNameLocal,  # RepresentativeNameLocal
            f"{data.birthyear}-{data.birthmonth:02d}-{data.birthday:02d}",  # RepresentativeBirthDay
            data.presidentAddress,  # RepresentativeAddress
            data.presidentAddressLocal,  # RepresentativeAddressLocal
            data.purpose1,  # BusinessPurpose1
            data.purpose2,  # BusinessPurpose2
            data.purpose3,  # BusinessPurpose3
            data.purpose4,  # BusinessPurpose4
            data.purpose5,  # BusinessPurpose5
            data.email,  # Email Address
            folder_url or "",  # FolderURL（空文字も許容）
        ]
        
        # Spreadsheetに追加（ヘッダー行の下に追加）
        body = {
            'values': [row_data]
        }
        
        result = sheets_service.spreadsheets().values().append(
            spreadsheetId=GOOGLE_SHEETS_SPREADSHEET_ID,
            range='A2',  # ヘッダー行（A1）の下から開始
            valueInputOption='RAW',
            insertDataOption='INSERT_ROWS',
            body=body
        ).execute()
        
        print(f"✅ Spreadsheet記録成功: {result.get('updates').get('updatedCells')} セル更新")
    except Exception as e:
        print(f"❌ Spreadsheet記録エラー: {e}")
        _log_google_exception("append_to_spreadsheet", e)
        # エラーが発生しても処理を続行


# 新テンプレート（temporary フォルダ）のパス
_TEMPORARY_DIR = os.path.normpath(os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "temporary"))
REGISTRATION_TEMPLATE_NEW = os.path.join(_TEMPORARY_DIR, "★❷(A商号)_登記申請_20250827_テンプレート_v1.0.docx")

# 登記申請
@app.post("/generate-registration-application")
async def generate_registration_application(data: FormData, folder_id: str | None = None):
    if os.path.exists(REGISTRATION_TEMPLATE_NEW):
        template_path = REGISTRATION_TEMPLATE_NEW
    else:
        template_path = "template-word-registration-application.docx"
        if not os.path.exists(template_path) and TEMPLATE_DIR:
            template_path = os.path.join(TEMPLATE_DIR, template_path)

    if not os.path.exists(template_path):
        raise HTTPException(status_code=500, detail="Template file not found")
    try:
        doc = Document(template_path)
    except Exception as e:
        doc = Document(os.path.join(TEMPLATE_DIR, template_path) if TEMPLATE_DIR else template_path)

    output_path = generate_output_filename(data.companyName, "registration-application.docx")

    # 登記申請で置換する対象: (A商号), (A商号のフリガナ), (C社員住所), (D社員氏名・カタカナ), (D社員氏名・英語)
    # (A商号のフリガナ) と (D社員氏名・カタカナ) は固有名詞として自然なカタカナ表記に変換。住所もできるだけ自然なカタカナにする。
    translated_company_name = await company_name_to_katakana(data.companyName)
    katakana_president_name = await name_to_katakana_roman_only(data.presidentName)
    address_katakana = await get_address_katakana_for_documents(data.presidentAddress)

    replacements = {
        "(A商号)": data.companyName,
        "(A商号のフリガナ)": translated_company_name,
        "(C社員住所)": address_katakana,
        "(D社員氏名・カタカナ)": katakana_president_name,
        "(D社員氏名・英語)": data.presidentName,
    }
    replace_in_docx_keeping_style(doc, replacements)

    # 生成された Word ファイルを保存
    print(f"✅ Wordファイルを保存: {output_path}")
    doc.save(output_path)

    # Googleドライブにアップロード
    try:
        target_folder_id = folder_id or GOOGLE_DRIVE_FOLDER_ID
        if not target_folder_id:
            print("⚠️ GOOGLE_DRIVE_FOLDER_IDが設定されていません。ファイルはローカルにのみ保存されます。")
        else:
            file_name = os.path.basename(output_path)
            file_id = await upload_file_to_drive(output_path, target_folder_id, file_name)
            print(f"✅ Googleドライブにアップロード完了: {file_id}")
    except Exception as e:
        print(f"⚠️ Googleドライブアップロードエラー（処理は続行）: {e}")

    return {"message": "Registration application generated", "filename": output_path}

    # headers = {
    #     "Content-Disposition": "attachment; filename=created_registration.docx",
    #     "Content-Type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    # }

    # with open(output_path, "rb") as file:
    #     return Response(content=file.read(), headers=headers, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

@app.get("/download-registration-application")
def download_registration_application(filename: str):
    if not filename:
        raise HTTPException(status_code=400, detail="Filename is required")
    
    file_path = filename
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found")

    try:
        with open(file_path, "rb") as file:
            file_stream = BytesIO(file.read())
    except Exception as e:
        with open(os.path.join(TEMPLATE_DIR, file_path), "rb") as file:
            file_stream = BytesIO(file.read())

    headers = {
        "Content-Disposition": f"attachment; filename={filename}",
        "Content-Type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    }
    return Response(content=file_stream.getvalue(), headers=headers, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# 定款（新テンプレート: temporary/★❶(A商号)_定款_20250827_テンプレート_v1.0.docx を優先）
ARTICLE_TEMPLATE_NEW = os.path.join(_TEMPORARY_DIR, "★❶(A商号)_定款_20250827_テンプレート_v1.0.docx")

# 定款作成
@app.post("/generate-article-of-incorporation")
async def generate_article_of_incorporation(data: FormData, folder_id: str | None = None):
    if os.path.exists(ARTICLE_TEMPLATE_NEW):
        template_path = ARTICLE_TEMPLATE_NEW
    else:
        template_path = "template-word-article-of-incorporation.docx"
        if not os.path.exists(template_path) and TEMPLATE_DIR:
            template_path = os.path.join(TEMPLATE_DIR, template_path)

    if not os.path.exists(template_path):
        raise HTTPException(status_code=500, detail="Template file not found")
    try:
        doc = Document(template_path)
    except Exception as e:
        doc = Document(os.path.join(TEMPLATE_DIR, template_path) if TEMPLATE_DIR else template_path)

    output_path = generate_output_filename(data.companyName, "article-of-incorporation.docx")

    # 定款で置換する対象: (A商号), (B目的1)~(B目的5), (C社員住所), (D社員氏名・カタカナ)。
    # (D社員氏名・カタカナ) と住所は固有名詞として自然なカタカナ表記に変換する。
    katakana_president_name = await name_to_katakana_roman_only(data.presidentName)
    address_katakana = await get_address_katakana_for_documents(data.presidentAddress)
    translated_purpose = await translate_text(data.purpose1)
    translated_purpose2 = await translate_text(data.purpose2)
    translated_purpose3 = await translate_text(data.purpose3)
    translated_purpose4 = await translate_text(data.purpose4)
    translated_purpose5 = await translate_text(data.purpose5)

    replacements = {
        "(A商号)": data.companyName,
        "(B目的1)": translated_purpose,
        "(B目的2)": translated_purpose2,
        "(B目的3)": translated_purpose3,
        "(B目的4)": translated_purpose4,
        "(B目的5)": translated_purpose5,
        "(C社員住所)": address_katakana,
        "(D社員氏名・カタカナ)": katakana_president_name,
    }
    replace_in_docx_keeping_style(doc, replacements)

    # 生成された Word ファイルを保存
    doc.save(output_path)
    
    # Googleドライブにアップロード
    try:
        target_folder_id = folder_id or GOOGLE_DRIVE_FOLDER_ID
        if not target_folder_id:
            print("⚠️ GOOGLE_DRIVE_FOLDER_IDが設定されていません。ファイルはローカルにのみ保存されます。")
        else:
            file_name = os.path.basename(output_path)
            file_id = await upload_file_to_drive(output_path, target_folder_id, file_name)
            print(f"✅ Googleドライブにアップロード完了: {file_id}")
    except Exception as e:
        print(f"⚠️ Googleドライブアップロードエラー（処理は続行）: {e}")

    return {"message": "Article of incorporation generated", "filename": output_path}

    # headers = {
    #     "Content-Disposition": "attachment; filename=created_incorparticles.docx",
    #     "Content-Type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    # }

    # with open(output_path, "rb") as file:
    #     return Response(content=file.read(), headers=headers, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

@app.get("/download-article-of-incorporation")
def download_article_of_incorporation(filename: str):
    if not filename:
        raise HTTPException(status_code=400, detail="Filename is required")
    
    file_path = filename
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found")

    with open(file_path, "rb") as file:
        file_stream = BytesIO(file.read())

    headers = {
        "Content-Disposition": f"attachment; filename={filename}",
        "Content-Type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    }
    return Response(content=file_stream.getvalue(), headers=headers, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# 印鑑届出書（新テンプレート: temporary/★❸(A商号)_印鑑届書_20250827_テンプレート_v1.0.xlsx を優先）
SEAL_TEMPLATE_NEW = os.path.join(_TEMPORARY_DIR, "★❸(A商号)_印鑑届書_20250827_テンプレート_v1.0.xlsx")

@app.post("/generate-seal-registration")
async def generate_seal_registration(data: FormData, folder_id: str | None = None):
    if os.path.exists(SEAL_TEMPLATE_NEW):
        template_path = SEAL_TEMPLATE_NEW
    else:
        template_path = "template-excel-seal-registration.xlsx"
        if not os.path.exists(template_path) and TEMPLATE_DIR:
            template_path = os.path.join(TEMPLATE_DIR, "template-excel-seal-registration.xlsx")

    if not os.path.exists(template_path):
        raise HTTPException(status_code=500, detail="Excel template file not found")

    wb = load_workbook(template_path)
    ws = wb.active

    # 氏名・(A商号のフリガナ)・住所は固有名詞として自然なカタカナ表記に変換する。
    katakana_president_name = await name_to_katakana_roman_only(data.presidentName)
    address_katakana = await get_address_katakana_for_documents(data.presidentAddress)
    birth_str = str(data.birthyear) + "年" + str(data.birthmonth) + "月" + str(data.birthday) + "日"

    if template_path == SEAL_TEMPLATE_NEW:
        # 印鑑届書で置換する対象: (A商号), (A商号のフリガナ), (C社員住所), (D社員氏名・カタカナ), (D社員氏名・英語), (E社員生年月日)
        # (A商号のフリガナ) と (D社員氏名・カタカナ) は固有名詞として自然なカタカナ表記に変換する。
        translated_company_name = await company_name_to_katakana(data.companyName)
        replacements = {
            "(A商号)": data.companyName,
            "(A商号のフリガナ)": translated_company_name,
            "(C社員住所)": address_katakana,
            "(D社員氏名・カタカナ)": katakana_president_name,
            "(D社員氏名・英語)": data.presidentName,
            "(E社員生年月日)": birth_str,
        }
        for sheet in wb.worksheets:
            for row in sheet.iter_rows():
                for cell in row:
                    if cell.value and isinstance(cell.value, str):
                        for ph, value in replacements.items():
                            if ph in cell.value:
                                cell.value = cell.value.replace(ph, value)
    else:
        # 旧テンプレート: 固定セルに設定
        def set_merged_cell_value(ws, cell_range, value):
            is_merged = False
            for merged_range in ws.merged_cells.ranges:
                if str(merged_range) == cell_range:
                    is_merged = True
                    break
            if is_merged:
                ws.unmerge_cells(cell_range)
            start_cell = cell_range.split(":")[0]
            ws[start_cell] = value
            if is_merged:
                ws.merge_cells(cell_range)

        set_merged_cell_value(ws, "AH7:BC9", data.companyName)
        set_merged_cell_value(ws, "P52:BC52", address_katakana)
        set_merged_cell_value(ws, "AH18:BC21", katakana_president_name)
        set_merged_cell_value(ws, "P53:BC53", katakana_president_name)
        set_merged_cell_value(ws, "AH22:BC24", birth_str)

    output_path = generate_output_filename(data.companyName, "seal-registration.xlsx")
    wb.save(output_path)

    # Googleドライブにアップロード
    try:
        target_folder_id = folder_id or GOOGLE_DRIVE_FOLDER_ID
        if not target_folder_id:
            print("⚠️ GOOGLE_DRIVE_FOLDER_IDが設定されていません。ファイルはローカルにのみ保存されます。")
        else:
            file_name = os.path.basename(output_path)
            file_id = await upload_file_to_drive(output_path, target_folder_id, file_name)
            print(f"✅ Googleドライブにアップロード完了: {file_id}")
    except Exception as e:
        print(f"⚠️ Googleドライブアップロードエラー（処理は続行）: {e}")

    return {"message": "Seal registration file generated", "filename": output_path}

    # headers = {
    #     "Content-Disposition": "attachment; filename=created_corporation_application.xlsx",
    #     "Content-Type": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    # }

    # with open(output_path, "rb") as file:
    #     return Response(content=file.read(), headers=headers, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

@app.get("/download-seal-registration")
def download_seal_registration(filename: str):
    if not filename:
        raise HTTPException(status_code=400, detail="Filename is required")
    
    file_path = filename
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found")

    with open(file_path, "rb") as file:
        file_stream = BytesIO(file.read())

    headers = {
        "Content-Disposition": f"attachment; filename={filename}",
        "Content-Type": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    }
    return Response(content=file_stream.getvalue(), headers=headers, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

# Spreadsheetに記録
@app.post("/record-to-spreadsheet")
def record_to_spreadsheet(data: FormData):
    """3つのファイル生成後にSpreadsheetに記録"""
    try:
        append_to_spreadsheet(data)
        return {"message": "Recorded to spreadsheet successfully"}
    except Exception as error:
        print(f"⚠️ Spreadsheet記録エラー: {error}")
        # エラーが発生しても処理を続行
        return {"message": "Spreadsheet recording failed but continuing", "error": str(error)}


async def _background_submit_task(data: FormData) -> None:
    """申請受付後のバックグラウンド処理: 3ファイル生成・Driveアップロード・Spreadsheet記録"""
    try:
        folder_id = None
        folder_url = ""
        try:
            if GOOGLE_DRIVE_FOLDER_ID:
                folder_id = await create_company_folder(data.companyName)
                if folder_id:
                    folder_url = f"https://drive.google.com/drive/folders/{folder_id}"
                    print(f"✅ 会社フォルダを使用: {folder_url}")
                else:
                    print("⚠️ 会社フォルダは作成できなかったため、親フォルダへ直接アップロードします。")
            else:
                print("⚠️ GOOGLE_DRIVE_FOLDER_ID が未設定のため、会社フォルダは作成しません。")
        except Exception as e:
            print(f"⚠️ 会社フォルダ作成エラー（処理は続行）: {e}")

        try:
            await generate_registration_application(data, folder_id=folder_id)
        except Exception as e:
            print(f"⚠️ 登記申請書生成エラー: {e}")
        try:
            await generate_article_of_incorporation(data, folder_id=folder_id)
        except Exception as e:
            print(f"⚠️ 定款生成エラー: {e}")
        try:
            await generate_seal_registration(data, folder_id=folder_id)
        except Exception as e:
            print(f"⚠️ 印鑑届出書生成エラー: {e}")
        try:
            append_to_spreadsheet(data, folder_url=folder_url)
        except Exception as e:
            print(f"⚠️ Spreadsheet記録エラー: {e}")
    except Exception as e:
        print(f"⚠️ バックグラウンド処理で予期せぬエラー: {e}")


@app.post("/submit-application")
async def submit_application(
    data: FormData,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """申請を受理し即座に200を返し、ファイル生成・Drive・Spreadsheetはバックグラウンドで実行（要ログイン・1ユーザー1申請）"""
    existing = db.query(Application).filter(Application.user_id == current_user.id).count()
    if existing > 0:
        raise HTTPException(status_code=409, detail="Application already submitted")
    now = datetime.now(timezone.utc)
    row = Application(
        user_id=current_user.id,
        payload=data.model_dump(),
        status=ApplicationStatus.PENDING.value,
        submitted_at=now,
        updated_at=now,
    )
    db.add(row)
    db.commit()
    asyncio.create_task(_background_submit_task(data))
    return {"message": "accepted"}


if __name__ == "__main__":
    port = int(os.getenv("PORT", "10000"))  # Render の環境変数から取得
    uvicorn.run(app, host="0.0.0.0", port=port)



