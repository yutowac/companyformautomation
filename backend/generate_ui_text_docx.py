"""UI 文言一覧 Word ドキュメントを生成（英語チェック用・説明は日本語）。"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

REPO_ROOT = Path(__file__).resolve().parent.parent
OUTPUT = REPO_ROOT / "docs" / "ui-text-inventory-en.docx"


def add_table(doc: Document, headers: tuple[str, ...], rows: list[tuple[str, ...]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
    for r_idx, row in enumerate(rows, start=1):
        cells = table.rows[r_idx].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
    doc.add_paragraph()


def main() -> None:
    doc = Document()
    title = doc.add_heading("画面文言一覧（英語チェック用）", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    intro = doc.add_paragraph()
    intro.add_run(
        "本ドキュメントは、アプリに表示される英語文言を画面ごとに整理したものです。"
        "「表示テキスト」列の英語表現をレビューしてください。\n"
        "出典: frontend/index.html, frontend/src/main.ts, frontend/src/locales/en.ts\n"
        "言語設定: 英語固定（i18n.ts）"
    ).font.size = Pt(10)

    doc.add_heading("共通（全画面）", level=1)
    add_table(
        doc,
        ("表示箇所", "表示テキスト（英語）", "説明"),
        [
            ("ブラウザタブ", "Company Registration Form", "ページタイトル（<title>）"),
            ("フッター", "© 2024 One-stop Inc. Japan. All rights reserved.", "全画面下部の著作権表示"),
        ],
    )

    doc.add_heading("1. ログイン画面（#/login）", level=1)
    add_table(
        doc,
        ("表示箇所", "表示テキスト（英語）", "説明"),
        [
            ("見出し", "Login", "ページタイトル"),
            ("説明文", "Sign in with your account credentials.", "ログインの案内文"),
            ("ラベル", "Email", "メールアドレス入力欄"),
            ("プレースホルダー", "you@example.com", "入力例"),
            ("ラベル", "Password", "パスワード入力欄"),
            ("ボタン", "Login", "ログイン実行"),
            ("エラー（動的）", "Please enter email and password.", "未入力時"),
            ("エラー（動的）", "Login failed", "ログイン失敗時のフォールバック"),
            ("エラー（動的）", "Incorrect email or password", "認証失敗時（API から）"),
        ],
    )

    doc.add_heading("2. ホーム画面（#/app/home）", level=1)
    add_table(
        doc,
        ("表示箇所", "表示テキスト（英語）", "説明"),
        [
            ("右上リンク", "→ Logout", "ログアウト"),
            ("見出し", "Home", "ページタイトル"),
            ("説明文", "Create a new application or check your application status.", "ホームの案内文"),
            ("カード見出し", "Application status", "申請状況エリアのラベル"),
            ("ステータス（動的）", "Loading...", "読み込み中の初期表示"),
            ("ステータス（動的）", "No application", "未申請のとき"),
            ("ステータス（動的）", "Pending", "申請中"),
            ("ステータス（動的）", "In review", "審査中"),
            ("ステータス（動的）", "Completed", "完了"),
            ("ステータス（動的）", "Rejected", "却下"),
            ("補足（動的）", "Submitted: {日時}", "申請日時（en-US 形式）"),
            ("ボタン", "New application", "新規申請へ遷移"),
            ("ボタン", "Check application status", "申請詳細へ遷移"),
        ],
    )
    doc.add_paragraph(
        "補足: 申請済みの場合「New application」は無効化、未申請の場合「Check application status」は無効化されます。"
        "ボタン文言自体は変わりません。"
    )

    doc.add_heading("3. 申請フォーム画面（#/app/form）", level=1)
    doc.add_heading("ヘッダー", level=2)
    add_table(
        doc,
        ("表示箇所", "表示テキスト（英語）", "説明"),
        [
            ("見出し", "One-stop Inc. Japan", "サービス名"),
            (
                "説明文",
                "Simply fill out the form below, and we will help you establish your company in Japan seamlessly.",
                "フォームの案内文",
            ),
        ],
    )

    doc.add_heading("会社名", level=2)
    add_table(
        doc,
        ("表示箇所", "表示テキスト（英語）", "説明"),
        [
            ("ラベル", "Company Name", "必須項目（* 表示あり）"),
            (
                "ツールチップ",
                "Please note: the legal suffix for the Company Name is filed as Godo Kaisha (LLC).",
                "「?」アイコンにホバーで表示",
            ),
            ("プレースホルダー", "YourCompany", "入力例"),
            ("固定表示", "LLC", "会社名入力欄の右側"),
        ],
    )

    doc.add_heading("代表者名", level=2)
    add_table(
        doc,
        ("表示箇所", "表示テキスト（英語）", "説明"),
        [
            ("セクション見出し", "Representative Name", "グループタイトル"),
            ("サブラベル", "English", "必須（* 表示あり）"),
            ("プレースホルダー", "John Smith", "英語名の入力例"),
            ("サブラベル", "Local Language", "任意入力"),
            ("プレースホルダー", "e.g., 張三 / محمد علي", "現地語名の入力例"),
        ],
    )

    doc.add_heading("代表者生年月日", level=2)
    add_table(
        doc,
        ("表示箇所", "表示テキスト（英語）", "説明"),
        [
            ("ラベル", "Representative Birth Day", "必須（* 表示あり）"),
            ("プレースホルダー", "yyyy / mm / dd", "年・月・日の各入力欄"),
        ],
    )

    doc.add_heading("代表者住所", level=2)
    add_table(
        doc,
        ("表示箇所", "表示テキスト（英語）", "説明"),
        [
            ("セクション見出し", "Representative Address", "グループタイトル"),
            ("サブラベル", "English", "必須（* 表示あり）"),
            ("プレースホルダー", "9-8-7 bef Apartment, Chiyoda, Tokyo", "英語住所の入力例"),
            ("サブラベル", "Local Language", "任意入力"),
            ("補足", "(Optional)", "現地語住所は任意である旨"),
            ("プレースホルダー", "e.g., 〒123-4567 東京 ... / ١٢٣ شارع ...", "現地語住所の入力例"),
        ],
    )

    doc.add_heading("事業目的", level=2)
    add_table(
        doc,
        ("表示箇所", "表示テキスト（英語）", "説明"),
        [
            ("ラベル", "Business Purpose 1", "必須（* 表示あり）"),
            ("プレースホルダー", "e.g., IT consulting business", "目的1の入力例"),
            ("ラベル", "Business Purpose 2 ～ 5", "「add」押下で追加表示"),
            ("プレースホルダー", "e.g., Software development business など", "目的2～5の入力例"),
            ("ボタン", "add", "事業目的フィールドを追加"),
            ("aria-label", "Add purpose", "追加ボタンのアクセシビリティ用ラベル"),
        ],
    )

    doc.add_heading("メール・操作ボタン", level=2)
    add_table(
        doc,
        ("表示箇所", "表示テキスト（英語）", "説明"),
        [
            ("ラベル", "Email Address", "必須（* 表示あり）"),
            ("プレースホルダー", "example@email.com", "メールの入力例"),
            ("ローディング", "Generating...", "送信中のスピナー横メッセージ"),
            ("ボタン", "Back to Home", "ホームへ戻る"),
            ("ボタン", "Done", "確認画面へ進む"),
        ],
    )

    doc.add_heading("ダウンロード（レガシー・通常は非表示）", level=2)
    add_table(
        doc,
        ("表示箇所", "表示テキスト（英語）", "説明"),
        [
            ("メッセージ", "Your documents are ready!! Please download them below.", "書類生成完了時"),
            ("ボタン", "Company Registration", "登記申請書のダウンロード"),
            ("ボタン", "Incorporation Articles", "定款のダウンロード"),
            ("ボタン", "Seal Registration", "印鑑届のダウンロード"),
        ],
    )

    doc.add_heading("4. 確認画面（#/app/confirm）", level=1)
    add_table(
        doc,
        ("表示箇所", "表示テキスト（英語）", "説明"),
        [
            ("見出し", "Confirm your information", "ページタイトル"),
            ("項目ラベル", "Company Name", "確認項目"),
            ("項目ラベル", "Representative Name (English)", "確認項目"),
            ("項目ラベル", "Representative Name (Local Language)", "確認項目"),
            ("項目ラベル", "Representative Birth Day", "確認項目"),
            ("項目ラベル", "Representative Address (English)", "確認項目"),
            ("項目ラベル", "Representative Address (Local Language)", "確認項目"),
            ("項目ラベル", "Business Purposes", "確認項目（複数はカンマ区切り）"),
            ("項目ラベル", "Email Address", "確認項目"),
            ("空欄時", "-", "値がないときの表示"),
            ("質問文", "Is this information correct?", "内容確認の問いかけ"),
            ("ボタン", "Edit", "フォームへ戻る"),
            ("ボタン", "Submit", "申請を送信"),
        ],
    )

    doc.add_heading("5. 完了画面（#/app/thanks）", level=1)
    add_table(
        doc,
        ("表示箇所", "表示テキスト（英語）", "説明"),
        [
            ("見出し", "Thank you for your application.", "送信完了メッセージ"),
            ("本文", "We will contact you once everything is ready.", "フォローアップ文"),
            ("ボタン", "Back to Home", "ホームへ戻る"),
        ],
    )

    doc.add_heading("6. 申請状況画面（#/app/applications）", level=1)
    add_table(
        doc,
        ("表示箇所", "表示テキスト（英語）", "説明"),
        [
            ("バッジ（動的）", "Pending / In review / Completed / Rejected", "左上のステータス表示"),
            ("バッジ（エラー時）", "Error", "読み込み失敗時"),
            ("見出し", "Application status", "ページタイトル"),
            ("補足（動的）", "Submitted: {日時}", "申請日時"),
            ("補足（エラー時）", "Failed to load application.", "読み込み失敗メッセージ"),
            ("項目ラベル", "確認画面と同じ8項目", "申請内容の詳細表示"),
            ("空欄時", "-", "値がないときの表示"),
            ("ボタン", "Back to Home", "ホームへ戻る"),
            ("テキスト", "Need to change your application?", "変更リクエストの前文"),
            ("リンク", "Click here", "Google フォームへ（別タブで開く）"),
        ],
    )

    doc.add_heading("7. アラート・エラーメッセージ（画面横断）", level=1)
    add_table(
        doc,
        ("表示タイミング", "表示テキスト（英語）", "説明"),
        [
            ("フォーム検証", "Please fill in all required fields.", "必須項目の未入力"),
            ("フォーム検証", "Please enter a valid email address.", "メール形式不正"),
            ("申請送信失敗", "Submission failed: {詳細}", "申請送信エラー"),
            ("未ログイン", "ログインが必要です", "日本語のまま — 英語化の候補"),
            ("二重申請", "Application already submitted", "API 409 応答"),
            ("ダウンロード失敗", "Download failed: {詳細}", "書類ダウンロードエラー"),
        ],
    )

    doc.add_heading("8. i18n に定義あるが画面に未使用の文言", level=1)
    doc.add_paragraph(
        "frontend/src/locales/en.ts に存在するが、現行 HTML フローでは表示されない文言です。"
    )
    add_table(
        doc,
        ("キー", "表示テキスト（英語）", "備考"),
        [
            ("page1Title", "Apply for incorporation with One-Stop inc", "旧フロー用"),
            ("page1Description", 'You will submit an application for incorporation. Click "Yes" to proceed to the form.', "旧フロー用"),
            ("page1Yes", "Yes", "旧フロー用"),
            ("submit", "Submit", "確認画面は HTML 直書きの Submit を使用"),
            ("presidentNameTooltip", "Please enter your name in English or your usual language.", "HTML に未接続"),
            ("presidentAddressTooltip", "Please enter your address in English or your usual language.", "HTML に未接続"),
            ("errorDownloadFailed", "Download failed", "api.ts では別文言を使用"),
        ],
    )

    doc.add_heading("英語チェック時の注意点", level=1)
    notes = [
        "大文字・小文字の統一: add（小文字）と Done / Submit / Login（先頭大文字）が混在しています。",
        "感嘆符: Your documents are ready!! は二重感嘆符です。",
        "ステータス表記: In review は文の途中のような表記（In Review ではない）。",
        "ページタイトル: Application status は status が小文字です。",
        "日本語残り: 申請送信時の「ログインが必要です」は英語化候補です。",
        "会社形態: フォームでは LLC、ツールチップでは Godo Kaisha (LLC) と説明が分かれています。",
    ]
    for note in notes:
        doc.add_paragraph(note, style="List Bullet")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Wrote: {OUTPUT}")


if __name__ == "__main__":
    main()
